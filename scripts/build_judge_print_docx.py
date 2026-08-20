from pathlib import Path
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "reports" / "judges" / "反诈项目价值与评测报告_评委版.docx"


def build(render_dir: Path):
    pages = sorted(render_dir.glob("page-*.png"), key=lambda p: int(p.stem.split("-")[-1]))
    if not pages:
        raise SystemExit(f"No page PNGs found in {render_dir}")
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0)
    doc.core_properties.title = "面向反诈教育的智能客服：项目价值与评测证据报告"
    doc.core_properties.subject = "评委审阅版（打印布局）"
    doc.core_properties.author = "项目团队"

    for index, page in enumerate(pages):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = 0
        paragraph.paragraph_format.space_after = 0
        paragraph.add_run().add_picture(str(page), width=Inches(7.5))
        if index < len(pages) - 1:
            doc.add_page_break()
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build(Path(sys.argv[1]))
