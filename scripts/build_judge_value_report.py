from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "reports" / "judges"
ASSET_DIR = OUT_DIR / "assets"
DOCX_PATH = OUT_DIR / "source" / "反诈项目价值与评测报告_编辑源.docx"
CHART_PATH = ASSET_DIR / "评测改进证据图.png"

FONT_LATIN = "Arial"
FONT_CJK = "Arial Unicode MS"
FONT_FILE = "/System/Library/Fonts/Hiragino Sans GB.ttc"

NAVY = "17324D"
BLUE = "2E74B5"
TEAL = "137A72"
GREEN = "237A57"
GOLD = "A26A00"
RED = "A13D3D"
INK = "17212B"
MUTED = "596773"
LIGHT = "F2F4F7"
PALE_BLUE = "EAF2F8"
PALE_GREEN = "E9F5F1"
PALE_GOLD = "FFF5DB"
WHITE = "FFFFFF"
GRID = "CBD5DE"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, size=10.5, bold=False, color=INK, italic=False, font=FONT_CJK):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = rgb(color)
    return run


def set_cell_fill(cell, color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def keep_row_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_text(cell, text, *, size=9.1, bold=False, color=INK, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    set_run_font(p.add_run(str(text)), size=size, bold=bold, color=color)


def add_table(doc, headers, rows, widths_dxa, *, header_fill=LIGHT, font_size=9.1):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        set_cell_fill(table.rows[0].cells[idx], header_fill)
        set_cell_text(table.rows[0].cells[idx], header, size=9.1, bold=True, color=NAVY)
    for row_values in rows:
        row = table.add_row()
        keep_row_together(row)
        for idx, value in enumerate(row_values):
            set_cell_text(row.cells[idx], value, size=font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_para(doc, text="", *, size=10.5, color=INK, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=6, line=1.10,
             keep_next=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.keep_with_next = keep_next
    set_run_font(p.add_run(text), size=size, color=color, bold=bold, italic=italic)
    return p


def add_rich_para(doc, parts, *, size=10.5, after=6, line=1.10, fill=None, border=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    for text, options in parts:
        set_run_font(p.add_run(text), size=options.get("size", size), bold=options.get("bold", False),
                     color=options.get("color", INK), italic=options.get("italic", False))
    if fill:
        p_pr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        p_pr.append(shd)
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.right_indent = Inches(0.12)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(8)
    if border:
        p_pr = p._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "6")
        bottom.set(qn("w:color"), border)
        p_bdr.append(bottom)
        p_pr.append(p_bdr)
    return p


def add_bullet(doc, text, *, color=INK):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    set_run_font(p.add_run(text), size=10.3, color=color)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    set_run_font(p.add_run(text), size=10.3)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    set_run_font(p.add_run(text), size={1: 16, 2: 13, 3: 12}[level], bold=True,
                 color=BLUE if level < 3 else NAVY)
    return p


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_run_font(run, size=8.5, color=MUTED)


def draw_chart():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    width, height = 1680, 800
    image = Image.new("RGB", (width, height), f"#{WHITE}")
    draw = ImageDraw.Draw(image)
    title_font = ImageFont.truetype(FONT_FILE, 46, index=0)
    label_font = ImageFont.truetype(FONT_FILE, 29, index=0)
    small_font = ImageFont.truetype(FONT_FILE, 24, index=0)
    value_font = ImageFont.truetype(FONT_FILE, 32, index=0)

    draw.text((70, 38), "两条回归证据线均显示问题被持续收敛", font=title_font, fill=f"#{NAVY}")

    def panel(x0, y0, x1, y1, title):
        draw.rounded_rectangle((x0, y0, x1, y1), radius=12, fill="#F7F9FB", outline=f"#{GRID}", width=2)
        draw.text((x0 + 28, y0 + 22), title, font=label_font, fill=f"#{NAVY}")

    panel(70, 120, 820, 730, "400 条内部回归：Top-1 迭代")
    vals = [50.75, 72.75, 92.50, 98.00]
    names = ["初始", "接口优化", "行为规则 v2", "行为规则 v3"]
    colors = ["9BAAB6", "5C92BC", "247C75", "237A57"]
    base_y, max_h = 610, 380
    for i, (v, name, color) in enumerate(zip(vals, names, colors)):
        x = 115 + i * 170
        bar_h = int(v / 100 * max_h)
        draw.rounded_rectangle((x, base_y - bar_h, x + 105, base_y), radius=8, fill=f"#{color}")
        draw.text((x + 52, base_y - bar_h - 42), f"{v:.2f}%", font=value_font, fill=f"#{NAVY}", anchor="mm")
        draw.multiline_text((x + 52, base_y + 24), name, font=small_font, fill=f"#{MUTED}", anchor="ma", align="center", spacing=3)

    panel(860, 120, 1610, 730, "210 条困难集：本轮优化前后")
    metrics = [("风险类型 Top-1", 56.84, 100.0, True), ("正常对照误报", 90.0, 0.0, False)]
    for row, (name, before, after, higher_better) in enumerate(metrics):
        y = 250 + row * 230
        draw.text((905, y - 62), name, font=label_font, fill=f"#{NAVY}")
        draw.text((905, y), "优化前", font=small_font, fill=f"#{MUTED}")
        draw.rectangle((1010, y + 2, 1010 + int(before * 4.6), y + 40), fill="#B7C1C9")
        draw.text((1515, y + 21), f"{before:.2f}%", font=value_font, fill=f"#{RED if not higher_better else MUTED}", anchor="rm")
        draw.text((905, y + 72), "优化后", font=small_font, fill=f"#{MUTED}")
        effective = after if higher_better else 100 - after
        draw.rectangle((1010, y + 74, 1010 + int(effective * 4.6), y + 112), fill=f"#{GREEN}")
        draw.text((1455, y + 93), f"{after:.2f}%", font=value_font, fill="#FFFFFF", anchor="rm")
    draw.text((870, 685), "注：210 条为开发回归集，不是独立盲测。", font=small_font, fill=f"#{MUTED}")
    image.save(CHART_PATH, quality=95)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_CJK
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for level, size, before, after, color in ((1, 16, 16, 8, BLUE), (2, 13, 12, 6, BLUE), (3, 12, 8, 4, NAVY)):
        style = styles[f"Heading {level}"]
        style.font.name = FONT_CJK
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = FONT_CJK
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run("面向反诈教育的智能客服"), size=8.5, bold=True, color=NAVY)
    set_run_font(p.add_run("  |  项目价值与评测证据"), size=8.5, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    set_run_font(p.add_run("评委审阅版  ·  第 "), size=8.5, color=MUTED)
    add_page_number(p)
    set_run_font(p.add_run(" 页"), size=8.5, color=MUTED)


def audit_docx(path: Path):
    with ZipFile(path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        styles_xml = archive.read("word/styles.xml").decode("utf-8")
        assert "w:pgSz w:w=\"12240\" w:h=\"15840\"" in document_xml
        assert "w:top=\"1440\"" in document_xml
        assert "w:tblW w:type=\"dxa\" w:w=\"9360\"" in document_xml
        assert "w:line=\"264\"" in styles_xml or "w:line=\"264.0\"" in styles_xml
        assert "Arial Unicode MS" in styles_xml


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    draw_chart()
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    doc.core_properties.title = "面向反诈教育的智能客服：项目价值与评测证据报告"
    doc.core_properties.subject = "评委审阅版"
    doc.core_properties.author = "项目团队"

    add_para(doc, "项目价值与评测证据报告", size=11, bold=True, color=TEAL, after=10)
    add_para(doc, "面向反诈教育的\n智能客服", size=27, bold=True, color=NAVY, after=8, line=1.02)
    add_para(doc, "从风险识别到可执行劝阻：一套可复核、可迭代的反诈交互系统", size=13, color=MUTED, after=14)
    add_rich_para(doc, [
        ("报告结论  ", {"bold": True, "color": WHITE, "size": 11}),
        ("项目已经形成“识别风险—解释证据—给出处置动作”的工程闭环；本轮困难开发集问题已闭环，但仍需独立盲测验证泛化能力。", {"color": WHITE, "size": 11}),
    ], fill=NAVY, after=12, line=1.18)

    add_table(doc, ["证据项", "当前结果", "评委应如何理解"], [
        ("210 条困难开发集", "Top-1 100%\n正常对照误报 0%", "证明本轮暴露的边界问题已被工程化修复，不等同于盲测成绩"),
        ("完整 HTTP 链路", "210/210 成功\n10 并发，P95 76.83 ms", "证明结果不只来自离线函数，真实接口链路可稳定返回"),
        ("400 条内部回归", "Top-1 50.75% → 98.00%", "证明优化过程可量化、可复现，低覆盖类型持续收敛"),
        ("规则可靠性", "18/18 单元测试通过\nwarning 样本 0", "证明关键组合、冲突矩阵、多轮状态与动作元数据均有回归保护"),
    ], [2100, 2200, 5060], header_fill=PALE_BLUE, font_size=8.9)
    add_para(doc, "报告日期：2026 年 8 月 2 日    |    评测口径：本地确定性回归 + HTTP 端到端压测", size=9, color=MUTED, after=2)
    add_para(doc, "证据边界：210 条样本由 105 条种子及其语义改写构成，共享案件族，明确归入开发集。", size=9, color=RED, bold=True, after=0)

    doc.add_page_break()
    add_heading(doc, "一、项目解决的不是“答题”，而是风险处置", 1)
    add_para(doc, "多数反诈产品只回答“这像不像诈骗”。真实用户更需要的是：在口语化、信息不完整、类型相互重叠的描述中，系统能否识别当前风险，说明为什么，并给出此刻最该做的动作。本项目把这三件事放进同一条交互链路。")
    add_rich_para(doc, [
        ("用户描述  →  场景路由  →  行为组合提取  →  风险类型/阶段判断  →  止损与核验动作", {"bold": True, "color": NAVY, "size": 11.2}),
    ], fill=PALE_BLUE, after=10)

    add_heading(doc, "1. 识别价值：处理真实口语和相邻类型", 2)
    add_para(doc, "系统不再只依赖“刷单”“客服”“游戏交易”等类型名称，而是组合判断接触渠道、交易对象、付款方式、账号控制、验证码、远程控制、承诺收益和退款条件等行为。这样可以处理用户没有说出诈骗名称、同一句话同时含有多个风险线索的情况。")
    add_heading(doc, "2. 解释价值：让判断可审计", 2)
    add_para(doc, "规则引擎输出主类型、候选类型、风险分、命中特征、风险阶段和干预动作。评委可以追溯“为什么判成这一类”，研发也能定位混淆来源，而不是只能接受一个不可解释的标签。")
    add_heading(doc, "3. 行动价值：把识别转化为止损", 2)
    add_para(doc, "对正在发生的风险，回答优先停止转账、停止共享屏幕或停止提供验证码，再引导官方渠道核验、保存证据、账户保护和报警止付；对正常官方流程，则明确提示当前未见明显诈骗特征，避免制造恐慌。")

    add_heading(doc, "二、核心技术改进", 1)
    add_table(doc, ["此前问题", "本轮改进", "带来的实际效果"], [
        ("类型关键词互相覆盖", "以多项行为组合设置优先级与冲突矩阵", "游戏交易、屏幕共享、投资等主要混淆对被拆开"),
        ("未知特征导致整条规则跳过", "补齐特征注册并校验加载结果", "210 条 HTTP 评测 warning 样本降为 0"),
        ("正常交易也被高风险提示", "增加官方渠道、合同流程、否定表达和安全支付识别", "困难集正常对照误报由 90% 降为 0%"),
        ("对话层用原文反推诈骗类型", "以风险引擎明确结果为准，未知/正常不再强行补类型", "减少“底层判正常、上层又报诈骗”的链路冲突"),
        ("风险回答缺少场景动作", "依据验证码、远程控制、培训费等目标选择干预动作", "回答从通用提醒升级为场景化处置"),
    ], [2100, 3450, 3810], font_size=8.8)

    doc.add_page_break()
    add_heading(doc, "三、量化证据：持续收敛，而非一次性“调到满分”", 1)
    doc.add_picture(str(CHART_PATH), width=Inches(6.35))
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    add_para(doc, "图 1  两条开发回归证据线。400 条数据来自项目案例衍生；210 条数据来自人工口语化种子及语义改写。两者均用于开发回归，不作为独立盲测。", size=8.5, color=MUTED, after=8)

    add_heading(doc, "1. 400 条内部回归：优化过程可量化", 2)
    add_table(doc, ["阶段", "Top-1", "Top-k", "风险接口 P95", "主要意义"], [
        ("初始规则", "50.75%", "68.25%", "90.9 ms", "暴露两卡、AI 换脸、票务、征信等低覆盖类型"),
        ("接口优化", "72.75%", "—", "—", "建立可运行的 HTTP 评测链路"),
        ("行为规则 v2", "92.50%", "97.75%", "237.9 ms", "行为组合与类型优先级显著减少混淆"),
        ("行为规则 v3", "98.00%", "98.50%", "102.6 ms", "剩余错误收敛到游戏交易等少数边界"),
    ], [1620, 1050, 1050, 1500, 4140], font_size=8.7)
    add_para(doc, "来源：evaluation/reports/case_derived_400_*.md。该数据与项目知识资产存在重叠，因此仅用于内部回归。", size=8.4, color=MUTED, after=8)

    add_heading(doc, "2. 210 条困难开发集：针对真实边界做定向验证", 2)
    add_para(doc, "这批样本集中加入游戏交易、客服/屏幕共享/验证码混淆、贷款/征信、正常交易、多轮过程以及招聘、退改签、租房、老师收费等长尾类型。优化前离线诊断为风险 Top-1 56.84%、正常对照误报 90%；优化后离线与 HTTP 链路均达到 Top-1 100%、正常对照误报 0%。")
    add_rich_para(doc, [
        ("关键解释：", {"bold": True, "color": RED}),
        ("该结果证明“已知问题已闭环”，不证明未知真实分布上的泛化准确率为 100%。", {"bold": True, "color": INK}),
    ], fill=PALE_GOLD, after=10)

    doc.add_page_break()
    add_heading(doc, "四、完整 HTTP 压测结果", 1)
    add_para(doc, "本轮已实际启动本地服务，并通过真实 HTTP 风险接口发送全部 210 条请求，而不是直接调用规则函数。压测使用 10 个并发工作线程，验证了路由、请求解析、风险服务、规则引擎和响应序列化的完整链路。")
    add_table(doc, ["指标", "结果", "判读"], [
        ("总请求", "210", "覆盖困难开发集全部样本"),
        ("成功请求", "210 / 210", "接口错误 0"),
        ("并发度", "10", "已验证小规模并发，不代表 50/100 并发容量"),
        ("平均延迟", "55.65 ms", "包含完整本地 HTTP 往返"),
        ("P50 / P95 / 最大值", "43.82 / 76.83 / 290.30 ms", "尾延迟可控，仍需更高并发下复测"),
        ("风险 Top-1 / Top-k", "100% / 100%", "190 条风险样本"),
        ("正常对照误报", "0%", "20 条正常对照"),
        ("warning 样本", "0", "未知特征不再导致规则整条跳过"),
    ], [2050, 2600, 4710], header_fill=PALE_GREEN, font_size=9.0)
    add_para(doc, "来源：evaluation/reports/user_augmented_210_http_evaluation.json；生成时间 2026-08-02 18:54:25（Asia/Shanghai）。", size=8.4, color=MUTED, after=10)

    add_heading(doc, "离线与 HTTP 为什么都要测", 2)
    add_bullet(doc, "离线评测定位规则本身的分类与误报问题，单条 P95 为 4.11 ms。")
    add_bullet(doc, "HTTP 评测覆盖真实服务链路，单条 P95 为 76.83 ms，可发现接口错误、序列化问题和运行时 warning。")
    add_bullet(doc, "两套结果一致，说明本轮修复没有只停留在测试函数中。")

    add_heading(doc, "五、代表性问题如何被修复", 1)
    add_table(doc, ["困难场景", "原先容易混淆", "现在使用的行为组合", "期望输出"], [
        ("游戏账号被找回/装备交易", "未知、客服、中奖", "游戏资产 + 账号/装备交付 + 找回/拒付/保证金", "游戏交易诈骗 + 平台申诉/留证"),
        ("客服要求远程或验证码", "统一判为冒充客服", "客服身份 + 远程控制/屏幕共享，或验证码 + 扣款", "按实际控制手段优先分类并给对应止损动作"),
        ("网恋对象拉投资", "普通投资或熟人借钱", "情感关系 + 投资引导 + 资金投入/无法提现", "情感交友诱导投资"),
        ("官方渠道正常退款/交易", "高风险误报", "官方入口 + 原路退款/合同 + 无私下转账 + 否定风险动作", "正常/未知风险，不强行贴诈骗类型"),
    ], [1700, 1950, 3400, 2310], font_size=8.5)

    doc.add_page_break()
    add_heading(doc, "六、当前实现成熟度：哪些已经有，哪些还没有", 1)
    add_table(doc, ["能力", "状态", "已有证据或当前缺口"], [
        ("诈骗类型体系与知识资产", "已实现", "29 类诈骗类型、171 项特征、38 条风险规则；另有预防建议、案例、法条和官方来源等资产"),
        ("行为组合规则与冲突消解", "已实现", "覆盖游戏交易、屏幕共享、验证码、投资、征信、招聘、票务等相邻类别"),
        ("正常流程误报控制", "已实现（开发集）", "20 条困难正常对照误报 0%；需要更大盲测集确认"),
        ("结构化风险卡与动作建议", "已实现", "输出类型、分数、阶段、特征、候选类型及场景化干预动作"),
        ("完整 HTTP 回归压测", "已实现（10 并发）", "210 请求零错误，P95 76.83 ms"),
        ("关键规则单元测试", "已实现", "18/18 通过，覆盖冲突矩阵、多轮状态、否定表达和干预元数据"),
        ("独立冻结盲测", "尚未完成", "当前 210/400 条均属于开发或内部回归，不可用于宣称真实泛化精度"),
        ("50/100 并发与长稳测试", "尚未完成", "当前只验证 10 并发；需补 P50/P95/P99、吞吐和资源占用"),
        ("线上 LLM 稳定性评测", "尚未完成", "需重复运行，报告均值、方差、超时和降级表现"),
        ("真实用户效果指标", "尚未完成", "需试点统计多轮完成率、劝阻动作采纳率、举报/止付成功率"),
        ("专项安全红队", "尚未完成", "需覆盖提示注入、隐私泄露、方言/ASR 错字和对抗表达"),
    ], [2500, 1800, 5060], font_size=8.6)

    add_heading(doc, "七、为什么这个项目有参赛价值", 1)
    add_number(doc, "问题真实：样本来自用户自然表达，包含犹豫、情绪、错别字、信息缺失和跨类型线索，贴近真实求助而非标准题干。")
    add_number(doc, "方案完整：不是单一分类器，而是从路由、特征、风险阶段到干预动作的一条可运行产品链路。")
    add_number(doc, "改进可证明：400 条回归 Top-1 从 50.75% 提升到 98.00%；本轮困难集的低准确与高误报均被量化闭环。")
    add_number(doc, "结果可解释：确定性规则、稳定类型 ID、候选类型和命中特征便于审计，也便于在安全场景中快速修正。")
    add_number(doc, "态度可信：项目明确承认开发集边界，并给出独立盲测、容量测试和真实效果验证计划。")

    doc.add_page_break()
    add_heading(doc, "八、下一阶段验证计划", 1)
    add_table(doc, ["优先级", "验证任务", "验收指标"], [
        ("P0", "冻结 500–1000 条独立双人盲标测试集；按案件族切分", "报告 Top-1、Macro-F1、每类召回、混淆矩阵和正常误报率"),
        ("P0", "增加正常交易、模糊表达、多轮状态变化", "正常误报率、风险召回与分阶段动作准确率"),
        ("P1", "10/50/100 并发阶梯压测与 30 分钟长稳", "吞吐、P50/P95/P99、错误率、CPU/内存"),
        ("P1", "在线 LLM 模式重复测试及确定性降级", "均值/方差、超时率、模板回退成功率"),
        ("P1", "小规模真实用户试点", "对话完成率、关键动作采纳率、满意度与误导投诉"),
        ("P2", "安全与鲁棒性红队", "提示注入、隐私泄露、ASR 错字、方言和对抗改写通过率"),
    ], [1050, 4700, 3610], font_size=8.8)

    add_heading(doc, "九、结论", 1)
    add_rich_para(doc, [
        ("可以向评委证明的价值是：", {"bold": True, "color": NAVY, "size": 11}),
        ("项目已经把反诈知识转化为一套可运行、可解释、可压测、可持续回归的风险处置系统，并对最突出的分类混淆和正常误报问题完成了工程闭环。", {"size": 11}),
    ], fill=PALE_GREEN, after=10, line=1.18)
    add_para(doc, "目前还不能证明的是：系统在全新真实分布上的准确率已经达到 100%，或在大规模并发和长期真实用户环境中已经稳定。把这部分作为下一阶段验证，而不是包装成现有成绩，反而使项目证据更可信。")

    add_heading(doc, "附录 A：证据索引与复现入口", 1)
    add_table(doc, ["证据", "项目内路径 / 命令"], [
        ("210 条 HTTP 结果", "evaluation/reports/user_augmented_210_http_evaluation.json"),
        ("210 条离线结果", "evaluation/reports/user_augmented_210_rule_evaluation.json"),
        ("400 条迭代结果", "evaluation/reports/case_derived_400_evaluation.md\nevaluation/reports/case_derived_400_http_p0_behavior_v3.md"),
        ("数据边界说明", "evaluation/annotations/user_annotation_manifest.json"),
        ("方向三基准", "reports/evaluation/direction3_evaluation.md"),
        ("复跑 210 条评测", "python scripts/run_user_annotation_evaluation.py --mode offline\npython scripts/run_user_annotation_evaluation.py --mode http --base-url http://127.0.0.1:8001 --workers 10"),
        ("复跑关键单测", "python -m unittest test.test_risk_engine -v"),
    ], [2500, 6860], font_size=8.2)
    add_para(doc, "审阅提示：所有百分比均保留两位小数；延迟为本机本地服务测量，不能直接外推到生产部署。", size=8.5, color=MUTED, after=0)

    doc.save(DOCX_PATH)
    audit_docx(DOCX_PATH)
    print(DOCX_PATH)
    print(CHART_PATH)


if __name__ == "__main__":
    build()
